"""
文檔基本設定

頁面尺寸改為 A4 (21cm x 29.7cm)
邊距都設定為 2.5 cm
添加了頁首（顯示報告標題和日期）和頁尾（顯示公司名稱和頁碼）
頁碼格式更改為「第 X 頁」


字體與文字格式

區分中文（微軟正黑體）和英文（Arial）字體
標題 H1 設為 16pt 深藍色 (#003366)
標題 H2 設為 14pt 灰藍色 (#336699)
標題 H3 設為 12pt 黑色
正文設為 11pt


段落設定

行距設為 1.5 行
段前後間距設為 6 pt
段落對齊方式為兩端對齊
首行縮排設為 0.75 cm（約2字元）


列表格式

無序列表使用圓點，有序列表使用阿拉伯數字
多層級列表格式：1., 1.1, 1.1.1
列表間距與段落相同


圖片處理

最大寬度限制為 15 cm
圖片置中內嵌
圖片說明置於下方，置中對齊，字體 10pt


表格格式

邊框設為 0.5pt 灰色
表頭背景色為 #F2F2F2，粗體，置中對齊
表格內容靠左對齊


特殊元素處理

超連結：保留下劃線，字體顏色改為黑色
引用區塊：左邊縮排 1 cm，使用斜體字
程式碼區塊：使用 Consolas 10pt，灰色背景


其他格式要求

添加完整封面頁
自動生成包含 H1 至 H3 的目錄
文檔背景色為白色


匯出相關

檔案格式為 .docx（Office 2010+ 相容）
設定文檔屬性（作者、標題、公司）
"""

from exporters.base import Exporter
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import io
import re
import requests
import base64
from PIL import Image
from datetime import datetime
from utils.logger import get_logger

logger = get_logger("docx_exporter")

class DocxExporter(Exporter):
    """
    專門處理 Docx (.docx) 檔案匯出的類別。
    根據預定義的格式設定轉換HTML內容至Docx檔案。
    支援Word圖表佔位符處理和自動目錄生成。
    """

    def __init__(self, 
                 content_dict: dict, 
                 company_info: dict, 
                 charts_data: dict = None, 
                 charts_position_info: dict = None
        ):
        super().__init__(content_dict, company_info)
        
        # 保存圖表數據
        self.charts_data = charts_data or {}
        self.charts_position_info = charts_position_info or {}
        
        # 定義顏色
        self.colors = {
            'deep_blue': RGBColor(0, 51, 102),  # #003366
            'gray_blue': RGBColor(51, 102, 153),  # #336699
            'black': RGBColor(0, 0, 0),
            'gray': RGBColor(204, 204, 204),  # #CCCCCC
            'light_gray': RGBColor(242, 242, 242),  # #F2F2F2
            'code_bg': RGBColor(244, 244, 244)  # #F4F4F4
        }
        
        # 定義字體設定
        self.fonts = {
            'title': {'name_zh': '微軟正黑體', 'name_en': 'Arial', 'size': Pt(16), 'bold': True, 'color': self.colors['deep_blue']},
            'h1': {'name_zh': '微軟正黑體', 'name_en': 'Arial', 'size': Pt(16), 'bold': True, 'color': self.colors['deep_blue']},
            'h2': {'name_zh': '微軟正黑體', 'name_en': 'Arial', 'size': Pt(14), 'bold': True, 'color': self.colors['gray_blue']},
            'h3': {'name_zh': '微軟正黑體', 'name_en': 'Arial', 'size': Pt(12), 'bold': True, 'color': self.colors['black']},
            'normal': {'name_zh': '微軟正黑體', 'name_en': 'Arial', 'size': Pt(11), 'bold': False, 'color': self.colors['black']},
            'caption': {'name_zh': '微軟正黑體', 'name_en': 'Arial', 'size': Pt(10), 'bold': False, 'color': self.colors['black']},
            'header_footer': {'name_zh': '微軟正黑體', 'name_en': 'Arial', 'size': Pt(9), 'bold': False, 'color': self.colors['black']},
            'code': {'name': 'Consolas', 'size': Pt(10), 'bold': False, 'color': self.colors['black']}
        }
        
        # 定義段落設定
        self.paragraph_settings = {
            'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'space_before': Pt(6),
            'space_after': Pt(6),
            'line_spacing': 1.5,
            'first_line_indent': Cm(0.75)  # 約2字元
        }

    @staticmethod
    def enable_field_update_on_open(document):
        """
        把 <w:updateFields w:val="true"/> 寫進 word/settings.xml，
        讓 Word 每次開檔自動重算 TOC／頁碼／交叉參照。
        """
        settings = document.settings.element
        # 若已經存在就不再重複塞
        if settings.find(qn('w:updateFields')) is None:
            update = OxmlElement('w:updateFields')
            update.set(qn('w:val'), 'true')
            settings.append(update)

    def apply_text_formatting(self, run, style_dict, is_english=False):
        """套用文字格式，根據中英文選擇不同字體"""
        if is_english and 'name_en' in style_dict:
            run.font.name = style_dict['name_en']
        elif 'name_zh' in style_dict:
            run.font.name = style_dict['name_zh']
        elif 'name' in style_dict:
            run.font.name = style_dict['name']
            
        if 'size' in style_dict:
            run.font.size = style_dict['size']
        if 'bold' in style_dict:
            run.font.bold = style_dict['bold']
        if 'italic' in style_dict:
            run.font.italic = style_dict['italic']
        if 'color' in style_dict:
            run.font.color.rgb = style_dict['color']
        if 'underline' in style_dict and style_dict['underline']:
            run.font.underline = True

    def apply_paragraph_settings(self, paragraph, settings=None):
        """套用段落格式"""
        if settings is None:
            settings = self.paragraph_settings
        
        paragraph.alignment = settings.get('alignment', WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.space_before = settings.get('space_before', Pt(0))
        paragraph.paragraph_format.space_after = settings.get('space_after', Pt(0))
        
        # 設定行距為1.5
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        
        paragraph.paragraph_format.first_line_indent = settings.get('first_line_indent', Cm(0))

    def setup_document(self, doc):
        """設定文檔基本屬性"""
        # 設定頁面尺寸和方向（A4大小）
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.orientation = WD_ORIENT.PORTRAIT
        
        # 設定頁面邊距
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
        # 設定頁首頁尾
        self.setup_header_footer(doc)
        
        # 設定文檔屬性
        doc.core_properties.title = self.title
        doc.core_properties.author = "產業研究報告系統"
        doc.core_properties.company = self.company_info.get("企業名稱", "")
        self.enable_field_update_on_open(doc)
    
    def setup_header_footer(self, doc):
        """設定頁首和頁尾"""
        section = doc.sections[0]
        
        # 設定頁首
        header = section.header
        header_table = header.add_table(1, 2, Cm(16))
        header_table.style = 'Table Grid'
        header_table.autofit = True
        
        # 左側：報告標題
        title_cell = header_table.cell(0, 0)
        title_p = title_cell.paragraphs[0]
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_p.add_run(self.title)
        self.apply_text_formatting(title_run, self.fonts['header_footer'])
        
        # 右側：日期
        date_cell = header_table.cell(0, 1)
        date_p = date_cell.paragraphs[0]
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_p.add_run(datetime.now().strftime('%Y/%m/%d'))
        self.apply_text_formatting(date_run, self.fonts['header_footer'])
        
        # 清除表格邊框
        for row in header_table.rows:
            for cell in row.cells:
                self.clear_cell_border(cell)
        
        # 設定頁尾
        footer = section.footer
        footer_table = footer.add_table(1, 2, Cm(16))
        footer_table.autofit = True
        
        # 左側：公司名稱或報告分類
        company_cell = footer_table.cell(0, 0)
        company_p = company_cell.paragraphs[0]
        company_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        company_name = self.company_info.get("企業名稱", "產業研究報告")
        company_run = company_p.add_run(company_name)
        self.apply_text_formatting(company_run, self.fonts['header_footer'])
        
        # 右側：頁碼
        page_cell = footer_table.cell(0, 1)
        page_p = page_cell.paragraphs[0]
        page_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        page_text = page_p.add_run("第 ")
        self.apply_text_formatting(page_text, self.fonts['header_footer'])
        
        # 添加頁碼域代碼
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        page_text._r.append(fldChar1)
        page_text._r.append(instrText)
        page_text._r.append(fldChar2)
        
        page_end = page_p.add_run(" 頁")
        self.apply_text_formatting(page_end, self.fonts['header_footer'])
        
        # 清除表格邊框
        for row in footer_table.rows:
            for cell in row.cells:
                self.clear_cell_border(cell)

    def create_cover_page(self, doc):
        """創建封面頁"""
        # 添加標題
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(self.title)
        self.apply_text_formatting(title_run, {'name_zh': '微軟正黑體', 'name_en': 'Arial', 'size': Pt(24), 'bold': True, 'color': self.colors['deep_blue']})
        
        # 添加空白行
        for _ in range(5):
            doc.add_paragraph()
        
        # 添加品牌和商品名稱
        brand_para = doc.add_paragraph()
        brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        brand_text = f"品牌: {self.company_info.get('品牌名稱', '')}"
        brand_run = brand_para.add_run(brand_text)
        self.apply_text_formatting(brand_run, self.fonts['h2'])
        
        product_para = doc.add_paragraph()
        product_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        product_text = f"商品: {self.company_info.get('商品名稱', '')}"
        product_run = product_para.add_run(product_text)
        self.apply_text_formatting(product_run, self.fonts['h2'])
        
        # 添加空白行
        for _ in range(5):
            doc.add_paragraph()
        
        # 添加日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_text = f"製作日期: {datetime.now().strftime('%Y年%m月%d日')}"
        date_run = date_para.add_run(date_text)
        self.apply_text_formatting(date_run, self.fonts['normal'])
        
        # 添加分頁符
        doc.add_page_break()

    def add_toc_with_proper_styles(self, doc):
        """添加完整的目錄，確保樣式正確生成"""
        try:
            # 添加目錄標題
            toc_heading = doc.add_heading("目錄", level=1)
            toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 修改目錄標題樣式
            for run in toc_heading.runs:
                self.apply_text_formatting(run, self.fonts['h1'])
            
            # 添加空白段落
            doc.add_paragraph()
            
            # 創建更可靠的TOC域代碼
            para = doc.add_paragraph()
            
            # 設置段落樣式
            para.style = 'Normal'
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 創建TOC的XML元素
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            # 使用更完整的TOC指令
            instrText.text = 'TOC \\o "1-3" \\h \\z \\t "Heading 1,1,Heading 2,2,Heading 3,3"'
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            
            # 添加默認的目錄內容提示
            t = OxmlElement('w:t')
            t.text = "請在Word中按F9鍵更新目錄"
            
            fldChar3 = OxmlElement('w:fldChar')  
            fldChar3.set(qn('w:fldCharType'), 'end')
            
            # 將所有元素添加到段落的run中
            run = para.add_run()
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(t)
            run._r.append(fldChar3)
            
            # 設置字體
            self.apply_text_formatting(run, self.fonts['normal'])
            
            # 添加目錄說明
            note_para = doc.add_paragraph()
            note_run = note_para.add_run("（註：開啟文件後請按F9鍵更新目錄）")
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = self.colors['gray_blue']
            note_run.italic = True
            note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加分頁符
            doc.add_page_break()
            
            logger.info("目錄已成功創建")
            
        except Exception as e:
            logger.error(f"創建目錄時發生錯誤: {e}")
            # 添加簡化的目錄
            self.add_simple_toc(doc)

    def add_simple_toc(self, doc):
        """添加簡化版目錄（備用方案）"""
        try:
            toc_heading = doc.add_heading("目錄", level=1)
            toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 手動創建目錄項目
            toc_items = []
            
            # 從content中提取標題
            for section, items in self.content.items():
                toc_items.append((section, 1))  # 主標題
                
                for subtitle in items.keys():
                    if subtitle != section:  # 避免重複
                        toc_items.append((subtitle, 2))  # 子標題
            
            # 添加目錄項目
            for title, level in toc_items:
                p = doc.add_paragraph()
                
                # 根據層級添加縮進
                if level == 1:
                    p.paragraph_format.left_indent = Cm(0)
                    run = p.add_run(f"• {title}")
                    self.apply_text_formatting(run, {'name_zh': '微軟正黑體', 'size': Pt(12), 'bold': True})
                else:
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run(f"○ {title}")
                    self.apply_text_formatting(run, {'name_zh': '微軟正黑體', 'size': Pt(11)})
                
                p.paragraph_format.space_after = Pt(3)
            
            doc.add_page_break()
            logger.info("簡化目錄已創建")
            
        except Exception as e:
            logger.error(f"創建簡化目錄失敗: {e}")

    def process_vanna_static_image(self, doc, img_static: bytes, title_text: str):
        """處理 Vanna 生成的靜態圖片 (bytes)"""
        try:
            if not img_static or len(img_static) == 0:
                logger.warning(f"圖表 {title_text} 的靜態圖片數據為空")
                self.add_error_placeholder(doc, f"圖表數據缺失: {title_text}")
                return
            
            # 將bytes轉換為BytesIO流
            img_stream = io.BytesIO(img_static)
            img_stream.seek(0)
            
            # 使用PIL驗證和處理圖片
            with Image.open(img_stream) as pil_img:
                # 檢查圖片格式和尺寸
                format_name = pil_img.format
                width, height = pil_img.size
                mode = pil_img.mode
                
                logger.info(f"Vanna圖表 - 格式: {format_name}, 尺寸: {width}x{height}, 模式: {mode}")
                
                # 如果圖片是 RGBA 或其他模式，轉換為 RGB
                if mode in ('RGBA', 'LA', 'P'):
                    logger.info(f"轉換圖片模式從 {mode} 到 RGB")
                    pil_img = pil_img.convert('RGB')
                
                # 計算適合的尺寸
                max_width_cm = 15  # 最大寬度 15cm
                max_width_px = max_width_cm * 37.795  # 1cm ≈ 37.795 pixels at 96 DPI
                
                if width > max_width_px:
                    # 計算縮放比例
                    scale = max_width_px / width
                    new_width = int(width * scale)
                    new_height = int(height * scale)
                    logger.info(f"Vanna圖表將被縮放到: {new_width}x{new_height}")
                    
                    # 重設大小
                    pil_img = pil_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # 將處理後的圖片保存到新的 BytesIO
                processed_stream = io.BytesIO()
                save_format = 'PNG' if format_name in ('PNG', 'GIF') else 'JPEG'
                pil_img.save(processed_stream, format=save_format, quality=85, optimize=True)
                processed_stream.seek(0)
                
                logger.info(f"Vanna圖表預處理完成，保存格式: {save_format}")
            
            # 計算Word中的顯示尺寸
            processed_stream.seek(0)
            with Image.open(processed_stream) as final_img:
                final_width, final_height = final_img.size
                
            # 轉換為 Word 的尺寸單位
            display_width = Cm(min(15, final_width * 2.54 / 96))  # 96 DPI 轉 cm
            
            # 重置流位置並添加到文檔
            processed_stream.seek(0)
            
            # 添加圖片段落
            img_paragraph = doc.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 插入圖片
            run = img_paragraph.add_run()
            inline_shape = run.add_picture(processed_stream, width=display_width)
            
            logger.info(f"成功添加Vanna圖表到文檔: {title_text}")
            
            # 添加圖片說明
            if title_text and title_text.strip():
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 移除可能的後綴，如 "(發票數據)"
                clean_title = title_text.replace("(發票數據)", "").strip()
                caption_run = caption.add_run(f"圖: {clean_title}")
                self.apply_text_formatting(caption_run, self.fonts['caption'])
                logger.info(f"添加Vanna圖表說明: {clean_title}")
            
        except Exception as process_err:
            logger.error(f"Vanna圖表處理失敗: {process_err}")
            logger.error(f"錯誤詳情: {type(process_err).__name__}: {str(process_err)}")
            self.add_error_placeholder(doc, f"Vanna圖表處理失敗: {title_text}")

    def process_image(self, doc, img_element):
        """處理圖片元素，支援 base64 和 URL 圖片"""
        try:
            src = img_element.get('src')
            if not src:
                logger.warning("圖片缺少src屬性")
                return
            
            logger.info(f"開始處理圖片: {src[:100]}...")
            
            img_stream = None
            
            # 處理 base64 圖片
            if src.startswith('data:image/'):
                try:
                    logger.info("處理 base64 格式圖片")
                    # 解析 base64 數據
                    if ',' in src:
                        header, data = src.split(',', 1)
                        # 解碼 base64 數據
                        image_data = base64.b64decode(data)
                        img_stream = io.BytesIO(image_data)
                        logger.info(f"base64 圖片解碼成功，大小: {len(image_data)} bytes")
                    else:
                        raise ValueError("base64 格式不正確")
                        
                except Exception as b64_err:
                    logger.error(f"處理 base64 圖片失敗: {b64_err}")
                    self.add_error_placeholder(doc, f"base64 圖片處理失敗: {str(b64_err)}")
                    return
            
            # 處理 URL 圖片
            else:
                try:
                    logger.info(f"處理 URL 圖片: {src}")
                    # 處理相對URL (如果有基礎URL)
                    if hasattr(self, 'base_url') and not (src.startswith('http://') or src.startswith('https://')):
                        if src.startswith('/'):
                            src = f"{self.base_url.rstrip('/')}{src}"
                        else:
                            src = f"{self.base_url.rstrip('/')}/{src.lstrip('/')}"
                            
                    # 設置請求頭，模仿瀏覽器行為
                    headers = {
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                        'Accept': 'image/webp,image/apng,image/*,*/*;q=0.8',
                        'Accept-Language': 'zh-TW,zh;q=0.9,en;q=0.8',
                        'Accept-Encoding': 'gzip, deflate, br'
                    }
                    
                    # 下載圖片
                    response = requests.get(src, headers=headers, allow_redirects=True, timeout=15, stream=True)
                    response.raise_for_status()
                    
                    # 檢查內容類型
                    content_type = response.headers.get('Content-Type', '')
                    logger.info(f"圖片 Content-Type: {content_type}")
                    
                    if content_type and not content_type.startswith('image/'):
                        logger.warning(f"URL不是圖片: {src} (Content-Type: {content_type})")
                        self.add_error_placeholder(doc, f"非圖片內容: {src}")
                        return
                        
                    # 讀取圖片數據
                    img_stream = io.BytesIO(response.content)
                    logger.info(f"URL 圖片下載成功，大小: {len(response.content)} bytes")
                    
                except requests.exceptions.RequestException as req_err:
                    logger.warning(f"下載圖片 {src} 失敗: {req_err}")
                    self.add_error_placeholder(doc, f"圖片下載失敗: {src}")
                    return
                except Exception as url_err:
                    logger.error(f"處理 URL 圖片失敗: {url_err}")
                    self.add_error_placeholder(doc, f"URL 圖片處理失敗: {str(url_err)}")
                    return
            
            # 處理圖片數據並添加到文檔
            if img_stream:
                try:
                    # 先嘗試用 PIL 驗證圖片
                    img_stream.seek(0)
                    with Image.open(img_stream) as pil_img:
                        # 檢查圖片格式
                        format_name = pil_img.format
                        width, height = pil_img.size
                        mode = pil_img.mode
                        
                        logger.info(f"圖片資訊 - 格式: {format_name}, 尺寸: {width}x{height}, 模式: {mode}")
                        
                        # 如果圖片是 RGBA 或其他模式，轉換為 RGB
                        if mode in ('RGBA', 'LA', 'P'):
                            logger.info(f"轉換圖片模式從 {mode} 到 RGB")
                            pil_img = pil_img.convert('RGB')
                        
                        # 計算適合的尺寸
                        max_width_cm = 15  # 最大寬度 15cm
                        max_width_px = max_width_cm * 37.795  # 1cm ≈ 37.795 pixels at 96 DPI
                        
                        if width > max_width_px:
                            # 計算縮放比例
                            scale = max_width_px / width
                            new_width = int(width * scale)
                            new_height = int(height * scale)
                            logger.info(f"圖片將被縮放到: {new_width}x{new_height}")
                            
                            # 重設大小
                            pil_img = pil_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        
                        # 將處理後的圖片保存到新的 BytesIO
                        processed_stream = io.BytesIO()
                        save_format = 'PNG' if format_name in ('PNG', 'GIF') else 'JPEG'
                        pil_img.save(processed_stream, format=save_format, quality=85, optimize=True)
                        processed_stream.seek(0)
                        
                        logger.info(f"圖片預處理完成，保存格式: {save_format}")
                    
                    # 計算Word中的顯示尺寸
                    processed_stream.seek(0)
                    with Image.open(processed_stream) as final_img:
                        final_width, final_height = final_img.size
                        
                    # 轉換為 Word 的尺寸單位
                    display_width = Cm(min(15, final_width * 2.54 / 96))  # 96 DPI 轉 cm
                    
                    # 重置流位置並添加到文檔
                    processed_stream.seek(0)
                    
                    # 添加圖片段落
                    img_paragraph = doc.add_paragraph()
                    img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 插入圖片
                    run = img_paragraph.add_run()
                    inline_shape = run.add_picture(processed_stream, width=display_width)
                    
                    logger.info(f"成功添加圖片到文檔，顯示寬度: {display_width}")
                    
                    # 添加圖片說明（如果有alt屬性）
                    alt_text = img_element.get('alt', '').strip()
                    if alt_text:
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_run = caption.add_run(f"圖: {alt_text}")
                        self.apply_text_formatting(caption_run, self.fonts['caption'])
                        logger.info(f"添加圖片說明: {alt_text}")
                    
                except Exception as process_err:
                    logger.error(f"圖片處理或添加失敗: {process_err}")
                    logger.error(f"錯誤詳情: {type(process_err).__name__}: {str(process_err)}")
                    self.add_error_placeholder(doc, f"圖片處理失敗: {str(process_err)}")
                    
        except Exception as e:
            logger.error(f"處理圖片時發生未預期錯誤: {e}")
            logger.error(f"錯誤類型: {type(e).__name__}")
            self.add_error_placeholder(doc, f"圖片處理錯誤: {str(e)}")

    def add_error_placeholder(self, doc, error_message):
        """添加錯誤提示段落"""
        try:
            err_para = doc.add_paragraph()
            err_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            err_run = err_para.add_run(f"[{error_message}]")
            err_run.font.color.rgb = RGBColor(255, 0, 0)  # 紅色
            err_run.font.size = Pt(10)
            logger.info(f"添加錯誤提示: {error_message}")
        except Exception as placeholder_err:
            logger.error(f"無法添加錯誤提示: {placeholder_err}")

    def process_table(self, doc, table_element):
        """處理表格元素"""
        try:
            rows = table_element.find_all('tr')
            if not rows:
                return
                
            # 計算列數
            columns = max(len(row.find_all(['td', 'th'])) for row in rows)
            
            # 創建表格
            table = doc.add_table(rows=len(rows), cols=columns)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 設定表格邊框
            for row in table.rows:
                for cell in row.cells:
                    # 將 RGB 顏色值轉為十六進制
                    hex_color = '%02x%02x%02x' % (
                        self.colors['gray'][0], 
                        self.colors['gray'][1], 
                        self.colors['gray'][2]
                    )
                    self.set_cell_border(cell, border_type='single', border_size=2, border_color=hex_color)
            
            # 填充表格內容
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                is_header = i == 0 or row.find('th')
                
                for j, cell in enumerate(cells):
                    if j < columns:
                        text = cell.get_text().strip()
                        table_cell = table.cell(i, j)
                        
                        # 設定表頭樣式
                        if is_header:
                            cell_para = table_cell.paragraphs[0]
                            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = cell_para.add_run(text)
                            self.apply_text_formatting(run, {'name_zh': '微軟正黑體', 'name_en': 'Arial', 'size': Pt(11), 'bold': True})
                            
                            # 設定背景色
                            from docx.oxml.ns import nsdecls
                            from docx.oxml import parse_xml
                            shading_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                            table_cell._tc.get_or_add_tcPr().append(shading_xml)
                        else:
                            cell_para = table_cell.paragraphs[0]
                            cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = cell_para.add_run(text)
                            self.apply_text_formatting(run, self.fonts['normal'])
            
            # 設定表格自動適應內容
            table.autofit = True
            
        except Exception as e:
            logger.error(f"處理表格時出錯: {e}")

    def set_cell_border(self, cell, border_locations=None, border_type='single', border_size=2, border_color='000000'):
        """設置表格單元格的邊框"""
        if border_locations is None:
            border_locations = ['top', 'left', 'bottom', 'right']
        
        # 獲取 XML 元素的 tcPr (table cell properties)，如果不存在則創建
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # 添加 tcBorders 元素，如果不存在
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
        
        # 定義邊框的 XML 標籤
        border_types = {
            'top': 'w:top',
            'left': 'w:left',
            'bottom': 'w:bottom',
            'right': 'w:right'
        }
        
        # 為每個需要的位置設置邊框
        for border_location in border_locations:
            if border_location in border_types:
                tag = border_types[border_location]
                border_element = OxmlElement(tag)
                border_element.set(qn('w:val'), border_type)  # 邊框類型
                border_element.set(qn('w:sz'), str(border_size))  # 邊框大小
                border_element.set(qn('w:color'), border_color)  # 邊框顏色
                tcBorders.append(border_element)

    def clear_cell_border(self, cell, border_locations=None):
        """清除表格單元格的邊框"""
        if border_locations is None:
            border_locations = ['top', 'left', 'bottom', 'right']
        
        # 獲取 XML 元素的 tcPr (table cell properties)，如果不存在則創建
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # 添加 tcBorders 元素，如果不存在
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
        
        # 定義邊框的 XML 標籤
        border_types = {
            'top': 'w:top',
            'left': 'w:left',
            'bottom': 'w:bottom',
            'right': 'w:right'
        }
        
        # 為每個需要的位置設置邊框為 nil
        for border_location in border_locations:
            if border_location in border_types:
                tag = border_types[border_location]
                border_element = OxmlElement(tag)
                border_element.set(qn('w:val'), 'nil')  # 設置為 nil 表示無邊框
                tcBorders.append(border_element)

    def process_blockquote(self, doc, blockquote):
        """處理引用區塊"""
        try:
            text = blockquote.get_text().strip()
            if not text:
                return
                
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
            run = p.add_run(text)
            style = self.fonts['normal'].copy()
            style['italic'] = True
            self.apply_text_formatting(run, style)
            
        except Exception as e:
            logger.error(f"處理引用區塊時出錯: {e}")

    def process_code(self, doc, code_block):
        """處理程式碼區塊"""
        try:
            text = code_block.get_text().strip()
            if not text:
                return
                
            # 添加程式碼區塊
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
            # 使用等寬字體
            run = p.add_run(text)
            self.apply_text_formatting(run, self.fonts['code'])
            
            # 添加灰色背景和邊框（這需要直接操作XML）
            # 注意：python-docx不直接支持段落背景色，這裡只是簡單處理
            # 完整實現需要更複雜的XML操作
            run.font.highlight_color = 15  # wdGray15 (淺灰)
            
        except Exception as e:
            logger.error(f"處理程式碼區塊時出錯: {e}")

    def process_list(self, doc, list_element, is_ordered=False):
        """處理有序或無序列表"""
        list_items = list_element.find_all('li', recursive=False)

        for idx, li in enumerate(list_items, start=1):
            p = doc.add_paragraph()
            # 基本清單樣式
            if is_ordered:
                p.style = 'List Number'
            else:
                p.style = 'List Bullet'
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)

            # 有序清單手動加編號
            if is_ordered:
                prefix = f"{idx}. "
                num_run = p.add_run(prefix)
                self.apply_text_formatting(num_run, self.fonts['normal'])

            # 處理列表項內容
            for child in li.contents:
                if child.name in ('ul', 'ol'):   # 留給遞迴
                    continue
                self.process_text_with_formatting(p, child)

            # 巢狀清單遞迴
            nested_ul = li.find('ul', recursive=False)
            if nested_ul:
                self.process_list(doc, nested_ul, False)
            nested_ol = li.find('ol', recursive=False)
            if nested_ol:
                self.process_list(doc, nested_ol, True)

    def add_hyperlink(self, paragraph, run, url):
        """添加超連結到Word文檔中"""
        try:
            # 這個功能需要直接操作Word的XML
            # Define relationships in XML
            part = paragraph.part
            r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            
            # Get the run's XML element
            r = run._r
            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # The existing run for the text content
            hyperlink.append(r)
            
            # Add hyperlink to paragraph
            r_position = paragraph._p.index(r)
            paragraph._p[r_position] = hyperlink
            
        except Exception as e:
            logger.error(f"添加超連結時出錯: {e}")
            # 如果無法建立超連結，至少保留文字
            return

    def process_text_with_formatting(self, paragraph, element):
        """處理文字格式，包括粗體、斜體等，並保留超連結"""
        try:
            # 如果元素是純文字節點
            if isinstance(element, str):
                # 判斷是否為英文
                is_english = all(ord(c) < 128 for c in element.strip())
                run = paragraph.add_run(element)
                self.apply_text_formatting(run, self.fonts['normal'], is_english)
                return
                
            # 處理子元素
            for child in element.children:
                # 如果子元素是字符串（文字節點）
                if isinstance(child, str):
                    # 判斷是否為英文
                    is_english = all(ord(c) < 128 for c in child.strip())
                    run = paragraph.add_run(child)
                    self.apply_text_formatting(run, self.fonts['normal'], is_english)
                
                # 如果是加粗標籤
                elif child.name == 'strong' or child.name == 'b':
                    text = child.get_text()
                    is_english = all(ord(c) < 128 for c in text.strip())
                    run = paragraph.add_run(text)
                    style = self.fonts['normal'].copy()
                    style['bold'] = True
                    self.apply_text_formatting(run, style, is_english)
                
                # 如果是斜體標籤
                elif child.name == 'em' or child.name == 'i':
                    text = child.get_text()
                    is_english = all(ord(c) < 128 for c in text.strip())
                    run = paragraph.add_run(text)
                    style = self.fonts['normal'].copy()
                    style['italic'] = True
                    self.apply_text_formatting(run, style, is_english)
                
                # 如果是下劃線標籤
                elif child.name == 'u':
                    text = child.get_text()
                    is_english = all(ord(c) < 128 for c in text.strip())
                    run = paragraph.add_run(text)
                    style = self.fonts['normal'].copy()
                    style['underline'] = True
                    self.apply_text_formatting(run, style, is_english)
                
                # 如果是超連結
                elif child.name == 'a':
                    text = child.get_text()
                    href = child.get('href', '')
                    is_english = all(ord(c) < 128 for c in text.strip())
                    
                    # 添加超連結文字並設置格式
                    run = paragraph.add_run(text)
                    style = self.fonts['normal'].copy()
                    style['underline'] = True
                    self.apply_text_formatting(run, style, is_english)
                    
                    # 使用python-docx添加真正的超連結
                    if href:
                        self.add_hyperlink(paragraph, run, href)
                
                # 如果是其他帶格式的元素，遞歸處理
                elif hasattr(child, 'children'):
                    self.process_text_with_formatting(paragraph, child)
                    
        except Exception as e:
            logger.error(f"處理文字格式時出錯: {e}")
            run = paragraph.add_run(f"[格式處理錯誤: {str(e)}]")
            run.font.color.rgb = RGBColor(255, 0, 0)

    def process_plotly_chart(self, doc, plotly_div):
        """處理Plotly圖表元素"""
        try:
            # 查找相關的script標籤
            script_tag = plotly_div.find_next_sibling('script')
            if not script_tag:
                # 嘗試查找父元素下的script
                parent = plotly_div.parent
                if parent:
                    script_tag = parent.find('script', string=lambda text: text and 'Plotly.newPlot' in text)
            
            if not script_tag:
                logger.warning("找不到相關的Plotly腳本")
                self.add_error_placeholder(doc, "找不到圖表數據")
                return
            
            # 提取Plotly數據
            script_content = script_tag.string or script_tag.get_text()
            
            # 使用正則表達式提取Plotly.newPlot的參數
            import json
            plot_match = re.search(r'Plotly\.newPlot\s*\(\s*["\']([^"\']+)["\'],\s*(\[.*?\]),\s*(\{.*?\})', script_content, re.DOTALL)
            
            if not plot_match:
                logger.warning("無法解析Plotly數據")
                self.add_error_placeholder(doc, "無法解析圖表數據")
                return
            
            div_id, data_str, layout_str = plot_match.groups()
            
            try:
                # 解析JSON數據
                plot_data = json.loads(data_str)
                plot_layout = json.loads(layout_str)
                
                # 使用Plotly生成靜態圖片
                self.create_plotly_image(doc, plot_data, plot_layout)
                
            except json.JSONDecodeError as e:
                logger.error(f"JSON解析錯誤: {e}")
                self.add_error_placeholder(doc, "圖表數據格式錯誤")
                
        except Exception as e:
            logger.error(f"處理Plotly圖表時出錯: {e}")
            self.add_error_placeholder(doc, f"圖表處理失敗: {str(e)}")

    def create_plotly_image(self, doc, plot_data, plot_layout):
        """使用Plotly創建靜態圖片並插入Docx"""
        try:
            import plotly.graph_objects as go
            import plotly.io as pio
            from PIL import Image
            import io
            
            # 創建Plotly圖表
            fig = go.Figure(data=plot_data, layout=plot_layout)
            
            # 設置圖片尺寸和格式
            fig.update_layout(
                width=800,
                height=600,
                font=dict(size=12),
                margin=dict(l=50, r=50, t=50, b=50)
            )
            
            # 將圖表轉換為PNG圖片
            img_bytes = pio.to_image(fig, format="png")
            
            # 使用PIL處理圖片
            img_stream = io.BytesIO(img_bytes)
            
            with Image.open(img_stream) as pil_img:
                # 計算適合的尺寸
                width, height = pil_img.size
                max_width_cm = 15  # 最大寬度 15cm
                max_width_px = max_width_cm * 37.795  # 1cm ≈ 37.795 pixels at 96 DPI
                
                if width > max_width_px:
                    scale = max_width_px / width
                    new_width = int(width * scale)
                    new_height = int(height * scale)
                    pil_img = pil_img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    
                    # 保存調整後的圖片
                    processed_stream = io.BytesIO()
                    pil_img.save(processed_stream, format='PNG', quality=85, optimize=True)
                    processed_stream.seek(0)
                else:
                    img_stream.seek(0)
                    processed_stream = img_stream
            
            # 計算Word中的顯示尺寸
            display_width = Cm(min(15, width * 2.54 / 96))
            
            # 添加圖片到文檔
            img_paragraph = doc.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = img_paragraph.add_run()
            inline_shape = run.add_picture(processed_stream, width=display_width)
            
            logger.info(f"成功添加Plotly圖表到文檔")
            
            # 添加圖片說明
            title = plot_layout.get('title', {})
            if isinstance(title, dict):
                title_text = title.get('text', '')
            else:
                title_text = str(title)
                
            if title_text:
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(f"圖: {title_text}")
                self.apply_text_formatting(caption_run, self.fonts['caption'])
            
        except ImportError:
            logger.error("缺少必要的庫: plotly 和 kaleido")
            self.add_error_placeholder(doc, "缺少圖表處理庫，請安裝 plotly 和 kaleido")
        except Exception as e:
            logger.error(f"創建Plotly圖片時出錯: {e}")
            self.add_error_placeholder(doc, f"圖表生成失敗: {str(e)}")

    def process_html_content(self, doc, html_content):
        """處理HTML內容轉換為Word文檔 - 清理版本"""
        try:
            # 獲取前端傳來的位置信息
            charts_position_info = getattr(self, 'charts_position_info', {})
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 處理各種HTML元素
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'img', 'table', 'blockquote', 'pre', 'code', 'div']):
                if element.name == 'h1':
                    text = element.get_text().strip()
                    clean_text = re.sub(r'^\d+\.\s*\d+\.\s*', '', text)
                    
                    h1 = doc.add_heading('', level=1)
                    run = h1.add_run(clean_text)
                    self.apply_text_formatting(run, self.fonts['h1'])
                
                elif element.name == 'h2':
                    text = element.get_text().strip()
                    clean_text = re.sub(r'^\d+\.\s*\d+\.\s*', '', text)
                    
                    h2 = doc.add_heading('', level=2)
                    run = h2.add_run(clean_text)
                    self.apply_text_formatting(run, self.fonts['h2'])
                
                elif element.name == 'h3':
                    text = element.get_text().strip()
                    clean_text = re.sub(r'^\d+\.\s*\d+\.\s*', '', text)
                    
                    h3 = doc.add_heading('', level=3)
                    run = h3.add_run(clean_text)
                    self.apply_text_formatting(run, self.fonts['h3'])
                
                elif element.name == 'p':
                    # 跳過圖表錯誤信息和佔位符
                    text_content = element.get_text().strip()
                    if any(x in text_content for x in ['[圖表數據缺失:', '[找不到圖表:', '[圖表處理失敗:', '[WORD_CHART_']):
                        continue
                    
                    p = doc.add_paragraph()
                    self.apply_paragraph_settings(p)
                    self.process_text_with_formatting(p, element)
                
                elif element.name in ('ul', 'ol'):
                    self.process_list(doc, element, element.name == 'ol')

                elif element.name == 'img':
                    self.process_image(doc, element)
                
                elif element.name == 'div':
                    # 合併後的div元素處理邏輯
                    element_classes = element.get('class', [])
                    
                    if 'word-chart-placeholder' in element_classes:
                        # 處理圖表佔位符
                        placeholder_text = element.get_text().strip()
                        if placeholder_text.startswith('[WORD_CHART_') and placeholder_text.endswith(']'):
                            chart_id = placeholder_text[12:-1]  # 移除 [WORD_CHART_ 和 ]
                            
                            # 使用前端提供的位置信息插入圖表
                            self.insert_chart_with_position_info(doc, chart_id, charts_position_info)
                            logger.info(f"✅ 處理圖表佔位符: {chart_id}")
                        else:
                            logger.warning(f"圖表佔位符格式不正確: {placeholder_text}")
                            
                    elif 'plotly-graph-div' in element_classes:
                        # 處理Plotly圖表
                        self.process_plotly_chart(doc, element)
                            
                elif element.name == 'table':
                    self.process_table(doc, element)
                    
                elif element.name == 'blockquote':
                    self.process_blockquote(doc, element)
                    
                elif element.name in ('pre', 'code'):
                    self.process_code(doc, element)
                
            logger.info("HTML內容處理完成")
            
        except Exception as e:
            logger.error(f"處理HTML內容時出錯: {e}")
            doc.add_paragraph(f"處理內容時發生錯誤: {str(e)}")

    def insert_chart_with_position_info(self, doc, chart_id, position_info):
        """使用位置信息正確插入圖表 - 簡化版本"""
        try:
            # 從位置信息中找到對應的圖表
            target_chart = None
            for page_info in position_info.values():
                for chart_info in page_info:
                    if chart_info.get("chart_id") == chart_id:
                        target_chart = chart_info
                        break
                if target_chart:
                    break
            
            if target_chart:
                title_text = target_chart.get("title_text", "")
                
                # 統一只使用 img_static_b64 字段
                img_static_b64 = target_chart.get("img_static_b64")
                
                if img_static_b64:
                    try:
                        # 統一的解碼邏輯
                        img_static = base64.b64decode(img_static_b64)
                        self.process_vanna_static_image(doc, img_static, title_text)
                        
                        logger.info(f"✅ 成功插入圖表: {title_text} ({len(img_static)} bytes)")
                        logger.info(f"📍 位置: {target_chart.get('target_section', '未知章節')}")
                        
                    except Exception as decode_err:
                        logger.error(f"❌ 圖表 {chart_id} base64 解碼失敗: {decode_err}")
                        self.add_error_placeholder(doc, f"圖表解碼失敗: {chart_id}")
                else:
                    logger.warning(f"❌ 圖表 {chart_id} 缺少 img_static_b64 數據")
                    self.add_error_placeholder(doc, f"圖表數據缺失: {chart_id}")
            else:
                logger.warning(f"❌ 找不到圖表 {chart_id} 的位置信息")
                available_ids = [chart.get('chart_id') for page_charts in position_info.values() for chart in page_charts]
                logger.warning(f"可用圖表ID: {available_ids}")
                self.add_error_placeholder(doc, f"找不到圖表: {chart_id}")
            
        except Exception as e:
            logger.error(f"❌ 插入圖表 {chart_id} 時發生錯誤: {e}")
            self.add_error_placeholder(doc, f"圖表處理失敗: {str(e)}")
            
    def export(self, file_path: str):
        """改進匯出流程，確保圖表正確處理"""
        logger.info(f"開始匯出Word文件: {file_path}")
        
        try:
            # 創建新文檔
            doc = Document()
            
            # 設定文檔基本屬性
            self.setup_document(doc)
            
            # 創建封面頁
            self.create_cover_page(doc)
            
            # 使用增強版目錄創建
            self.add_toc_with_proper_styles(doc)
            
            # 詳細記錄圖表數據狀態
            if self.charts_data:
                logger.info(f"可用的圖表數據頁面: {list(self.charts_data.keys())}")
                total_charts = 0
                for page, charts in self.charts_data.items():
                    logger.info(f"頁面 '{page}' 有 {len(charts)} 個圖表")
                    total_charts += len(charts)
                    for i, chart in enumerate(charts):
                        chart_id = chart.get("chart_id", f"unknown_{i}")
                        title = chart.get("title_text", "無標題")
                        
                        # 檢查圖片數據狀態
                        has_img_static_b64 = "img_static_b64" in chart
                        has_img_static = "img_static" in chart
                        
                        img_size = 0
                        if has_img_static_b64:
                            try:
                                img_data = base64.b64decode(chart["img_static_b64"])
                                img_size = len(img_data)
                            except:
                                img_size = 0
                        elif has_img_static:
                            img_field = chart["img_static"]
                            if isinstance(img_field, dict) and "bytes" in img_field:
                                img_size = len(img_field["bytes"]) if img_field["bytes"] else 0
                            elif isinstance(img_field, (bytes, str)):
                                img_size = len(img_field) if img_field else 0
                        
                        logger.info(f"  圖表 {i}: ID={chart_id}, 標題={title}, "
                                  f"有b64={has_img_static_b64}, 有static={has_img_static}, "
                                  f"圖片大小={img_size} bytes")
                
                logger.info(f"總計 {total_charts} 個圖表")
            else:
                logger.warning("沒有圖表數據")
            
            # 處理各節內容
            for section_index, (section, items) in enumerate(self.content.items()):
                logger.info(f"處理章節: {section}")
                
                # 處理章節內容
                for subtitle, html_content in items.items():
                    if html_content:
                        logger.info(f"處理小節: {subtitle}")
                        
                        # 檢查並記錄圖表佔位符
                        if '[WORD_CHART_' in html_content:
                            import re
                            placeholders = re.findall(r'\[WORD_CHART_([^\]]+)\]', html_content)
                            logger.info(f"在 {subtitle} 中發現 {len(placeholders)} 個圖表佔位符: {placeholders}")
                        
                        # 使用新的處理邏輯
                        self.process_html_content(doc, html_content)
                
                # 添加分頁符，除非是最後一個章節
                if section_index < len(self.content) - 1:
                    doc.add_page_break()
            
            # 儲存文件
            doc.save(file_path)
            
            # 記錄完成信息
            total_charts = sum(len(charts) for charts in self.charts_data.values()) if self.charts_data else 0
            logger.info(f"Word文件已成功匯出至: {file_path}")
            logger.info(f"應該包含 {total_charts} 個圖表")
            logger.info(f"建議在Word中按F9鍵更新目錄")
            
        except Exception as e:
            logger.error(f"匯出Word文件時發生錯誤: {e}")
            logger.error(f"錯誤類型: {type(e).__name__}")
            
            # 創建詳細的錯誤報告
            try:
                error_doc = Document()
                error_doc.add_heading("Word匯出錯誤報告", level=0)
                error_doc.add_paragraph(f"匯出時間: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                error_doc.add_paragraph(f"目標文件: {file_path}")
                error_doc.add_paragraph(f"錯誤類型: {type(e).__name__}")
                error_doc.add_paragraph(f"錯誤信息: {str(e)}")
                
                # 添加調試信息
                error_doc.add_paragraph("調試信息:")
                error_doc.add_paragraph(f"- 章節數量: {len(self.content)}")
                error_doc.add_paragraph(f"- 圖表數據頁面: {list(self.charts_data.keys()) if self.charts_data else '無'}")
                
                total_charts = sum(len(charts) for charts in self.charts_data.values()) if self.charts_data else 0
                error_doc.add_paragraph(f"- 總圖表數量: {total_charts}")
                
                error_doc.add_paragraph("\n請將此報告發送給系統管理員。")
                error_doc.save(file_path)
                
                logger.info(f"錯誤報告已保存至: {file_path}")
                
            except Exception as report_error:
                logger.error(f"無法創建錯誤報告: {report_error}")